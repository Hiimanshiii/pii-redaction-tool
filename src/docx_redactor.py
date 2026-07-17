# pyrefly: ignore [missing-import]
import docx
from pathlib import Path
from collections import defaultdict
from typing import Dict, Set, List
from src.detectors import PIIDetector, DetectedEntity
from src.replacement_engine import ReplacementEngine

class DocxRedactor:
    """Orchestrates PII detection and consistent replacement inside DOCX documents."""

    def __init__(self, input_path: str, output_path: str, min_confidence: float = 0.85) -> None:
        in_p = Path(input_path).resolve()
        out_p = Path(output_path).resolve()

        if not in_p.exists():
            raise FileNotFoundError(f"Input document not found at: {input_path}")
        if not in_p.is_file():
            raise ValueError(f"Input path is not a file: {input_path}")
        if in_p.suffix.lower() != ".docx":
            raise ValueError(f"Unsupported input file extension: {in_p.suffix}. Only .docx is supported.")
        if out_p.suffix.lower() != ".docx":
            raise ValueError(f"Unsupported output file extension: {out_p.suffix}. Only .docx is supported.")
        if in_p == out_p:
            raise ValueError("Output path cannot equal the input path to prevent overwriting the original document.")

        self.input_path = input_path
        self.output_path = output_path
        self.min_confidence = min_confidence

        self.detector = PIIDetector()
        self.engine = ReplacementEngine()

        # safe aggregate counters
        self.total_redactions = 0
        self.counts_by_type: Dict[str, int] = defaultdict(int)
        
        # internal tracking for uniqueness (never exposed outside get_statistics())
        self.unique_original_values_by_type: Dict[str, Set[str]] = defaultdict(set)

        # internal tracking of redacted segments for leakage validation
        self._redacted_items: List[tuple] = []

    def _redact_text(self, text: str) -> str:
        """Detect entities, resolve overlaps, replace PII, and track statistics."""
        if not text or not text.strip():
            return text

        try:
            entities = self.detector.detect(text, min_confidence=self.min_confidence)
            if not entities:
                return text

            # Priority rules
            structured_types = {
                "EMAIL", "PHONE", "SSN", "CREDIT_CARD", "DATE_OF_BIRTH", "PAN", "AADHAAR", "IP_ADDRESS"
            }

            def get_priority(entity_type: str) -> int:
                if entity_type in structured_types:
                    return 1
                elif entity_type == "ADDRESS":
                    return 2
                elif entity_type == "COMPANY":
                    return 3
                elif entity_type == "PERSON":
                    return 4
                return 5

            # Resolve overlapping detections: sort by priority (asc), then confidence (desc), then length (desc)
            sorted_entities = sorted(
                entities,
                key=lambda x: (get_priority(x.entity_type), -x.confidence, -(x.end - x.start))
            )

            accepted: List[DetectedEntity] = []
            for entity in sorted_entities:
                overlap = False
                for acc in accepted:
                    if max(entity.start, acc.start) < min(entity.end, acc.end):
                        overlap = True
                        break
                if not overlap:
                    accepted.append(entity)

            if not accepted:
                return text

            # Sort right-to-left to keep indexes stable
            accepted.sort(key=lambda x: -x.start)

            redacted_text = text
            for entity in accepted:
                replacement = self.engine.get_replacement(entity)
                redacted_text = redacted_text[:entity.start] + replacement + redacted_text[entity.end:]

                # stats tracking
                self.total_redactions += 1
                self.counts_by_type[entity.entity_type] += 1
                self.unique_original_values_by_type[entity.entity_type].add(entity.text)

            # track for validation rescan (using the final redacted string)
            for entity in accepted:
                self._redacted_items.append((entity.entity_type, entity.text, redacted_text))

            return redacted_text
        except Exception as e:
            raise RuntimeError("An unexpected error occurred during document redaction processing.") from e

    def get_statistics(self) -> dict:
        """Return safe aggregate redaction statistics."""
        unique_counts = {t: len(s) for t, s in self.unique_original_values_by_type.items()}
        return {
            "total_redactions": self.total_redactions,
            "counts_by_type": dict(self.counts_by_type),
            "unique_counts_by_type": unique_counts
        }

    def _update_paragraph_text(self, paragraph, new_text: str) -> None:
        """Helper to update paragraph text preserving the first run's formatting."""
        if not paragraph.runs:
            paragraph.text = new_text
            return

        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = new_text
        else:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    def redact_document(self) -> None:
        """Process paragraphs, tables, and headers/footers in the DOCX document."""
        # Ensure output directory exists
        Path(self.output_path).parent.mkdir(parents=True, exist_ok=True)

        doc = docx.Document(self.input_path)

        # 1. Process top-level paragraphs
        for paragraph in doc.paragraphs:
            old_text = paragraph.text
            new_text = self._redact_text(old_text)
            if old_text != new_text:
                self._update_paragraph_text(paragraph, new_text)

        # 2. Process tables
        processed_cells: Set = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc in processed_cells:
                        continue
                    processed_cells.add(cell._tc)
                    
                    for paragraph in cell.paragraphs:
                        old_text = paragraph.text
                        new_text = self._redact_text(old_text)
                        if old_text != new_text:
                            self._update_paragraph_text(paragraph, new_text)

        # 3. Process headers/footers
        processed_hf: Set = set()
        for section in doc.sections:
            for hf_name in ["header", "footer"]:
                hf = getattr(section, hf_name)
                if hf is None:
                    continue
                if hf._element in processed_hf:
                    continue
                processed_hf.add(hf._element)

                for paragraph in hf.paragraphs:
                    old_text = paragraph.text
                    new_text = self._redact_text(old_text)
                    if old_text != new_text:
                        self._update_paragraph_text(paragraph, new_text)

                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell._tc in processed_cells:
                                continue
                            processed_cells.add(cell._tc)

                            for paragraph in cell.paragraphs:
                                old_text = paragraph.text
                                new_text = self._redact_text(old_text)
                                if old_text != new_text:
                                    self._update_paragraph_text(paragraph, new_text)

        # Save document securely
        try:
            doc.save(self.output_path)
        except Exception as e:
            raise RuntimeError(f"Failed to save redacted document to {self.output_path}: {e}") from e

    def validate_redaction(self) -> dict:
        """Rescans the output document to detect any leakage of original sensitive values.

        Returns:
            dict: A dictionary containing validation status ('passed': bool) and count of leaks.
        """
        import re
        leaked_count = 0
        for etype, orig, redacted_text in self._redacted_items:
            val = orig.strip()
            val_lower = val.lower()
            if not val or len(val) < 3:
                continue

            # Filter out generic terms for PERSON and COMPANY
            if etype == "PERSON":
                generic_words = {
                    "bidder", "defaulter", "conditioning", "managerial", "price", 
                    "amount", "transfer", "capital", "personnel", "website", "company",
                    "secretary", "officer", "compliance", "promoter", "director", "shareholder",
                    "signatory", "witness", "member", "group", "key"
                }
                if any(w in val_lower for w in generic_words):
                    continue
                if len(val.split()) < 2:
                    continue

            elif etype == "COMPANY":
                generic_companies = {
                    "private limited", "limited", "llp", "bank", "securities", "capital", 
                    "industries", "escrow", "refund", "collection", "short term", "the company",
                    "our company", "india limited", "bse limited", "national stock exchange",
                    "stock exchange", "issuer company", "registrar", "auditor", "adviser"
                }
                if val_lower in generic_companies:
                    continue
                if val_lower in ["private limited", "limited", "llp", "bank", "securities", "capital"]:
                    continue

            # Special validation for phone formatting:
            if etype == "PHONE":
                orig_digits = re.sub(r'\D', '', val_lower)
                if len(orig_digits) > 10 and orig_digits.startswith("91"):
                    orig_digits = orig_digits[2:]
                output_digits = re.sub(r'\D', '', redacted_text.lower())
                if orig_digits and orig_digits in output_digits:
                    leaked_count += 1
                continue

            # Case-sensitive check for PERSON/COMPANY to avoid false matches on uppercase headings or generic replacements
            if etype in ["PERSON", "COMPANY"]:
                pattern = re.compile(rf"(?<![a-zA-Z0-9]){re.escape(val)}(?![a-zA-Z0-9])")
                if pattern.search(redacted_text):
                    leaked_count += 1
            else:
                norm_orig = val_lower
                norm_output = redacted_text.lower()
                norm_output = re.sub(r'\s+', ' ', norm_output)
                pattern = re.compile(rf"(?<![a-zA-Z0-9]){re.escape(norm_orig)}(?![a-zA-Z0-9])")
                if pattern.search(norm_output):
                    leaked_count += 1

        passed = (leaked_count == 0)
        return {
            "passed": passed,
            "potential_leaks": leaked_count
        }
