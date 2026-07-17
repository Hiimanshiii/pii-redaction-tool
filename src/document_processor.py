from pathlib import Path
from typing import Dict, Any, List
# pyrefly: ignore [missing-import]
import docx

class DocumentProcessor:
    """Processes DOCX documents to extract statistics and text content."""

    def __init__(self, file_path: str) -> None:
        """Initialize the DocumentProcessor with a file path string.

        Args:
            file_path (str): Path to the DOCX file.
        """
        self.file_path = file_path
        self.document = None

    def load_document(self) -> None:
        """Load the DOCX document from the file path.

        Raises:
            FileNotFoundError: If the file does not exist at the specified path.
            ValueError: If the path is not a file, has an unsupported extension, or is corrupted.
        """
        path = Path(self.file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found at the path: {self.file_path}")
        if not path.is_file():
            raise ValueError(f"The path is not a file: {self.file_path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Unsupported file extension: {path.suffix}. Only .docx is supported.")
        
        try:
            self.document = docx.Document(path)
        except Exception as e:
            raise ValueError(f"Corrupted or unreadable DOCX document: {e}") from e

    def get_statistics(self) -> Dict[str, int]:
        """Calculate and return document statistics.

        Returns:
            Dict[str, int]: A dictionary containing paragraph and table metrics.
        
        Raises:
            ValueError: If load_document() has not been called.
        """
        if self.document is None:
            raise ValueError("Document is not loaded. Call load_document() first.")

        total_paragraphs = len(self.document.paragraphs)
        non_empty_paragraphs = sum(
            1 for p in self.document.paragraphs if p.text.strip()
        )
        
        total_tables = len(self.document.tables)
        total_table_rows = sum(len(table.rows) for table in self.document.tables)
        
        total_table_cells = 0
        for table in self.document.tables:
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc not in seen_cells:
                        seen_cells.add(cell._tc)
            total_table_cells += len(seen_cells)

        return {
            "total_paragraphs": total_paragraphs,
            "non_empty_paragraphs": non_empty_paragraphs,
            "total_tables": total_tables,
            "total_table_rows": total_table_rows,
            "total_table_cells": total_table_cells,
        }

    def extract_text(self) -> str:
        """Extract non-empty text from paragraphs and table cells.

        Deduplicates text from merged cells in tables.

        Returns:
            str: The combined extracted text, separated by newline characters.
        
        Raises:
            ValueError: If load_document() has not been called.
        """
        if self.document is None:
            raise ValueError("Document is not loaded. Call load_document() first.")

        extracted_lines = []

        # Extract text from paragraphs
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_lines.append(text)

        # Extract text from tables (deduplicating merged cells using XML element identity)
        for table in self.document.tables:
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc not in seen_cells:
                        seen_cells.add(cell._tc)
                        cell_text = cell.text.strip()
                        if cell_text:
                            extracted_lines.append(cell_text)

        return "\n".join(extracted_lines)

    def get_text_chunks(self) -> List[str]:
        """Extract each non-empty paragraph and unique table cell as a list of text chunks.

        Deduplicates table cells using XML element identity.

        Returns:
            List[str]: A list of text chunks.

        Raises:
            ValueError: If load_document() has not been called.
        """
        if self.document is None:
            raise ValueError("Document is not loaded. Call load_document() first.")

        chunks = []

        # Paragraph chunks
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        # Unique table cell chunks
        for table in self.document.tables:
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc not in seen_cells:
                        seen_cells.add(cell._tc)
                        cell_text = cell.text.strip()
                        if cell_text:
                            chunks.append(cell_text)

        return chunks

