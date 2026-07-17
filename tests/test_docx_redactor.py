import os
import unittest
# pyrefly: ignore [missing-import]
import docx
from src.docx_redactor import DocxRedactor

class TestDocxRedactor(unittest.TestCase):
    def setUp(self) -> None:
        self.input_file = "tests/temp_test_input.docx"
        self.output_file = "output/temp_test_output.docx"
        
        # Clean up existing files
        for f in [self.input_file, self.output_file]:
            if os.path.exists(f):
                try:
                    os.remove(f)
                except OSError:
                    pass
                
    def tearDown(self) -> None:
        for f in [self.input_file, self.output_file]:
            if os.path.exists(f):
                try:
                    os.remove(f)
                except OSError:
                    pass
                
    def test_docx_redaction_pipeline(self) -> None:
        # Create temp input document
        doc = docx.Document()
        
        # Add paragraphs
        doc.add_paragraph("Hello, my name is Sarthak Malvadkar and I work at KSH International Limited.")
        doc.add_paragraph("You can email me at sarthak@ksh.co.in or call +91 98765 43210.")
        doc.add_paragraph("No PII in this paragraph.")
        
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Sarthak Malvadkar in table cell."
        table.rows[0].cells[1].text = "Office No. 201, Tower 2, Kharadi, Pune, Maharashtra 411014"
        
        # Add merged cells
        cell_a = table.rows[1].cells[0]
        cell_b = table.rows[1].cells[1]
        cell_a.merge(cell_b)
        cell_a.text = "Merged cell containing sarthak@ksh.co.in"
        
        # Save temp input
        doc.save(self.input_file)
        
        # Redact
        redactor = DocxRedactor(
            input_path=self.input_file,
            output_path=self.output_file,
            min_confidence=0.85
        )
        redactor.redact_document()
        
        # Verify output exists
        self.assertTrue(os.path.exists(self.output_file))
        
        # Load redacted document
        redacted_doc = docx.Document(self.output_file)
        
        # Paragraphs check
        p0 = redacted_doc.paragraphs[0].text
        self.assertNotIn("Sarthak Malvadkar", p0)
        self.assertNotIn("KSH International Limited", p0)
        
        p1 = redacted_doc.paragraphs[1].text
        self.assertNotIn("sarthak@ksh.co.in", p1)
        self.assertNotIn("+91 98765 43210", p1)
        
        p2 = redacted_doc.paragraphs[2].text
        self.assertEqual(p2, "No PII in this paragraph.")
        
        # Table cells check
        c0 = redacted_doc.tables[0].rows[0].cells[0].text
        self.assertNotIn("Sarthak Malvadkar", c0)
        
        c1 = redacted_doc.tables[0].rows[0].cells[1].text
        self.assertNotIn("Kharadi", c1)
        
        # Merged cell check
        c_merged = redacted_doc.tables[0].rows[1].cells[0].text
        self.assertNotIn("sarthak@ksh.co.in", c_merged)
        
        # Check stats
        stats = redactor.get_statistics()
        self.assertGreater(stats["total_redactions"], 0)
        self.assertIn("PERSON", stats["counts_by_type"])
        self.assertIn("EMAIL", stats["counts_by_type"])
        self.assertNotIn("original_values", stats)  # Ensure no leak of raw data

    def test_same_person_consistent_replacement(self) -> None:
        doc = docx.Document()
        doc.add_paragraph("Sarthak Malvadkar is here.")
        doc.add_paragraph("Sarthak Malvadkar is there.")
        doc.save(self.input_file)

        redactor = DocxRedactor(self.input_file, self.output_file, min_confidence=0.85)
        redactor.redact_document()

        redacted_doc = docx.Document(self.output_file)
        name1 = redacted_doc.paragraphs[0].text.replace(" is here.", "")
        name2 = redacted_doc.paragraphs[1].text.replace(" is there.", "")

        self.assertEqual(name1, name2)
        self.assertNotEqual(name1, "Sarthak Malvadkar")

    def test_docx_redactor_validation_errors(self) -> None:
        # Non-existent file
        with self.assertRaises(FileNotFoundError):
            DocxRedactor("nonexistent_redactor_file.docx", self.output_file)

        # Create self.input_file for the other tests
        doc = docx.Document()
        doc.add_paragraph("Test text")
        doc.save(self.input_file)

        # Invalid output extension
        with self.assertRaises(ValueError):
            DocxRedactor(self.input_file, "out.txt")

        # Overwrite prevention
        with self.assertRaises(ValueError):
            DocxRedactor(self.input_file, self.input_file)

    def test_integration_redaction_leakage(self) -> None:
        import tempfile
        import os
        from docx import Document

        # Original PII test values (all synthetic and checksum-valid)
        original_pii = {
            "PERSON": "Aaditya Patel",
            "EMAIL": "aaditya.patel@example.com",
            "PHONE": "+91 98765 43210",
            "COMPANY": "Patel Enterprises Limited",
            "ADDRESS": "Flat No. 102, Shanti Sadan, Link Road, Mumbai 400001, Maharashtra, India",
            "PAN": "ABCDE1234F",
            "AADHAAR": "1234 5678 9010",
            "IP_ADDRESS": "192.0.2.1",
            "SSN": "123-45-6789",
            "CREDIT_CARD": "4111 1111 1111 1111",
            "DOB": "15 August 1980"
        }

        # Create a document containing these values
        doc = Document()
        doc.add_paragraph("This is a non-PII header sentence.")
        for pii_type, pii_val in original_pii.items():
            doc.add_paragraph(f"The user's {pii_type}: {pii_val}.")
        doc.add_paragraph("This is a non-PII footer sentence.")

        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, "input.docx")
            output_path = os.path.join(temp_dir, "output.docx")

            # Save original doc
            doc.save(input_path)

            # Get original size/modification time to verify it doesn't change
            orig_mtime = os.path.getmtime(input_path)
            orig_size = os.path.getsize(input_path)

            # Run DocxRedactor
            redactor = DocxRedactor(input_path, output_path, min_confidence=0.85)
            redactor.redact_document()

            # Verify input document was not modified
            self.assertEqual(os.path.getsize(input_path), orig_size)
            self.assertEqual(os.path.getmtime(input_path), orig_mtime)

            # Open redacted doc and extract text
            red_doc = Document(output_path)
            redacted_text = "\n".join(p.text for p in red_doc.paragraphs)

            # Verify no original PII remains
            for pii_type, pii_val in original_pii.items():
                self.assertNotIn(pii_val.lower(), redacted_text.lower())

            # Verify non-PII text remains
            self.assertIn("This is a non-PII header sentence.", redacted_text)
            self.assertIn("This is a non-PII footer sentence.", redacted_text)

            # Run post-redaction validation rescan on output
            validation = redactor.validate_redaction()
            self.assertTrue(validation["passed"])
            self.assertEqual(validation["potential_leaks"], 0)

    def test_validate_redaction_no_synthetic_leakage(self) -> None:
        import tempfile
        import os
        from docx import Document

        doc = Document()
        # Adding some PII that will be redacted
        doc.add_paragraph("Hello Sarthak Malvadkar and cs.connect@kshinternational.com")
        doc.save(self.input_file)

        redactor = DocxRedactor(self.input_file, self.output_file, min_confidence=0.85)
        redactor.redact_document()

        # Run validate_redaction
        validation = redactor.validate_redaction()
        # It must pass because original values are removed, and synthetic replacements are not leakage
        self.assertTrue(validation["passed"])
        self.assertEqual(validation["potential_leaks"], 0)

if __name__ == '__main__':
    unittest.main()
